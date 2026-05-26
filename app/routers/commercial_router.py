"""Commercial — clients (FF/Shipper), grilles tarifaires, offres, commandes.

Reprises de la V3.0.0 :
- Brackets dégressifs par volume (lt50→full ship).
- Génération de référence : ORD-YYYY-NNNN, RG-YYYY-NNNN, RO-YYYY-NNNN.
- Conversion offre → commande.
- Audit trail activity_logs sur toutes les actions.
"""
from __future__ import annotations

import io
from datetime import date as _date, datetime, timezone
from decimal import Decimal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from fastapi import APIRouter, Depends, Form, HTTPException, Request
from fastapi.responses import HTMLResponse, RedirectResponse, Response
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.database import get_db
from app.models.commercial import (
    CLIENT_TYPES, Client, Order, OrderAssignment,
    RateGrid, RateGridLine, RateOffer,
)
from app.models.leg import Leg
from app.permissions import require_permission
from app.services.activity import record as activity_record
from app.services.commercial import (
    bracket_rate, compute_offer_total, default_brackets_for,
    next_grid_reference, next_offer_reference, next_order_reference,
    pick_bracket,
)
from app.templating import templates

router = APIRouter(prefix="/commercial", tags=["commercial"])


# ────────────────────────────────────────────── Landing
@router.get("", response_class=HTMLResponse)
@router.get("/", response_class=HTMLResponse)
async def commercial_index(
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    clients_count = (await db.scalar(select(__import__("sqlalchemy").func.count(Client.id)))) or 0
    grids_active = (await db.scalar(
        select(__import__("sqlalchemy").func.count(RateGrid.id)).where(RateGrid.status == "active")
    )) or 0
    offers_open = (await db.scalar(
        select(__import__("sqlalchemy").func.count(RateOffer.id)).where(RateOffer.status.in_(("draft", "sent")))
    )) or 0
    orders_open = (await db.scalar(
        select(__import__("sqlalchemy").func.count(Order.id)).where(Order.status.in_(("draft", "confirmed", "loaded")))
    )) or 0
    last_orders = list((await db.execute(
        select(Order).options(selectinload(Order.client))
        .order_by(Order.created_at.desc()).limit(10)
    )).scalars().all())
    return templates.TemplateResponse(
        "staff/commercial/index.html",
        {
            "request": request, "user": user,
            "clients_count": clients_count,
            "grids_active": grids_active,
            "offers_open": offers_open,
            "orders_open": orders_open,
            "last_orders": last_orders,
        },
    )


# ────────────────────────────────────────────── Clients (FF / Shipper)
@router.get("/clients", response_class=HTMLResponse)
async def clients_list(
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    clients = list((await db.execute(
        select(Client).order_by(Client.name.asc())
    )).scalars().all())
    return templates.TemplateResponse(
        "staff/commercial/clients.html",
        {"request": request, "user": user, "clients": clients, "types": CLIENT_TYPES},
    )


@router.post("/clients")
async def client_create(
    request: Request,
    name: str = Form(...),
    client_type: str = Form(...),
    contact_name: str | None = Form(None),
    contact_email: str | None = Form(None),
    contact_phone: str | None = Form(None),
    country: str | None = Form(None),
    vat_number: str | None = Form(None),
    notes: str | None = Form(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    if client_type not in CLIENT_TYPES:
        raise HTTPException(status_code=400, detail="invalid client_type")
    c = Client(
        name=name.strip(),
        client_type=client_type,
        contact_name=(contact_name or "").strip() or None,
        contact_email=(contact_email or "").strip() or None,
        contact_phone=(contact_phone or "").strip() or None,
        country=(country or "").strip().upper()[:2] or None,
        vat_number=(vat_number or "").strip() or None,
        notes=notes,
    )
    db.add(c)
    await db.flush()
    await activity_record(
        db, action="create", user_id=user.id, user_name=user.full_name or user.username,
        user_role=user.role, module="commercial", entity_type="client",
        entity_id=c.id, entity_label=c.name,
        ip_address=_client_ip(request),
    )
    return RedirectResponse(url="/commercial/clients", status_code=303)


# ────────────────────────────────────────────── Rate grids
@router.get("/grids", response_class=HTMLResponse)
async def grids_list(
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    grids = list((await db.execute(
        select(RateGrid).options(selectinload(RateGrid.client), selectinload(RateGrid.lines))
        .order_by(RateGrid.created_at.desc())
    )).scalars().all())
    return templates.TemplateResponse(
        "staff/commercial/grids.html",
        {"request": request, "user": user, "grids": grids},
    )


@router.get("/grids/new", response_class=HTMLResponse)
async def grid_new_form(
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
) -> HTMLResponse:
    clients = list((await db.execute(
        select(Client).where(Client.is_active.is_(True)).order_by(Client.name)
    )).scalars().all())
    return templates.TemplateResponse(
        "staff/commercial/grid_form.html",
        {"request": request, "user": user, "clients": clients, "grid": None},
    )


@router.post("/grids")
async def grid_create(
    request: Request,
    client_id: int = Form(...),
    valid_from: str = Form(...),
    valid_to: str | None = Form(None),
    base_rate: float = Form(...),
    adjustment_index: float = Form(1.0),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    client = await db.get(Client, client_id)
    if client is None:
        raise HTTPException(status_code=404, detail="client introuvable")
    ref = await next_grid_reference(db)
    grid = RateGrid(
        reference=ref,
        client_id=client.id,
        status="draft",
        valid_from=_date.fromisoformat(valid_from),
        valid_to=_date.fromisoformat(valid_to) if valid_to else None,
        currency="EUR",
        base_rate_per_palette=Decimal(str(base_rate)),
        adjustment_index=Decimal(str(adjustment_index)),
    )
    db.add(grid)
    await db.flush()
    # Auto-create default brackets based on client_type
    for b in default_brackets_for(client.client_type):
        db.add(RateGridLine(
            grid_id=grid.id,
            bracket_key=b["key"],
            bracket_label=b["label"],
            max_qty=int(b["max_qty"]),
            coeff=Decimal(str(b["coeff"])),
        ))
    await db.flush()
    await activity_record(
        db, action="create", user_id=user.id, user_name=user.full_name or user.username,
        user_role=user.role, module="commercial", entity_type="rate_grid",
        entity_id=grid.id, entity_label=ref, ip_address=_client_ip(request),
    )
    return RedirectResponse(url=f"/commercial/grids/{grid.id}", status_code=303)


@router.get("/grids/{grid_id}", response_class=HTMLResponse)
async def grid_detail(
    grid_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    grid = (await db.execute(
        select(RateGrid).options(selectinload(RateGrid.client), selectinload(RateGrid.lines))
        .where(RateGrid.id == grid_id)
    )).scalar_one_or_none()
    if grid is None:
        raise HTTPException(status_code=404)
    return templates.TemplateResponse(
        "staff/commercial/grid_detail.html",
        {"request": request, "user": user, "grid": grid},
    )


@router.post("/grids/{grid_id}/activate")
async def grid_activate(
    grid_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    grid = await db.get(RateGrid, grid_id)
    if grid is None:
        raise HTTPException(status_code=404)
    grid.status = "active"
    await db.flush()
    await activity_record(
        db, action="update", user_id=user.id, user_name=user.full_name or user.username,
        user_role=user.role, module="commercial", entity_type="rate_grid",
        entity_id=grid.id, entity_label=grid.reference, detail="activated",
        ip_address=_client_ip(request),
    )
    return RedirectResponse(url=f"/commercial/grids/{grid_id}", status_code=303)


# ────────────────────────────────────────────── Offers
@router.get("/offers", response_class=HTMLResponse)
async def offers_list(
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    offers = list((await db.execute(
        select(RateOffer).order_by(RateOffer.created_at.desc())
    )).scalars().all())
    return templates.TemplateResponse(
        "staff/commercial/offers.html",
        {"request": request, "user": user, "offers": offers},
    )


@router.get("/offers/new", response_class=HTMLResponse)
async def offer_new_form(
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
) -> HTMLResponse:
    clients = list((await db.execute(
        select(Client).where(Client.is_active.is_(True)).order_by(Client.name)
    )).scalars().all())
    legs = list((await db.execute(
        select(Leg).order_by(Leg.etd.desc()).limit(50)
    )).scalars().all())
    grids = list((await db.execute(
        select(RateGrid).where(RateGrid.status == "active")
    )).scalars().all())
    return templates.TemplateResponse(
        "staff/commercial/offer_form.html",
        {"request": request, "user": user, "clients": clients,
         "legs": legs, "grids": grids, "offer": None},
    )


@router.post("/offers")
async def offer_create(
    request: Request,
    client_id: int = Form(...),
    title: str = Form(...),
    grid_id: int | None = Form(None),
    leg_id: int | None = Form(None),
    estimated_palettes: int = Form(0),
    valid_until: str | None = Form(None),
    notes: str | None = Form(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    if not await db.get(Client, client_id):
        raise HTTPException(status_code=404, detail="client introuvable")
    grid = await db.get(RateGrid, grid_id) if grid_id else None
    proposed_rate = Decimal("0")
    total = Decimal("0")
    if grid is not None and estimated_palettes > 0:
        # Build bracket lookup from lines
        lines = list((await db.execute(
            select(RateGridLine).where(RateGridLine.grid_id == grid.id)
            .order_by(RateGridLine.max_qty)
        )).scalars().all())
        if lines:
            picked = pick_bracket(
                [{"max_qty": l.max_qty, "coeff": float(l.coeff)} for l in lines],
                estimated_palettes,
            )
            if picked:
                proposed_rate = bracket_rate(
                    base_rate=grid.base_rate_per_palette,
                    coeff=picked["coeff"],
                    adjustment_index=grid.adjustment_index,
                )
                total = (proposed_rate * Decimal(estimated_palettes)).quantize(Decimal("0.01"))

    ref = await next_offer_reference(db)
    offer = RateOffer(
        reference=ref,
        client_id=client_id,
        grid_id=grid.id if grid else None,
        leg_id=leg_id,
        title=title.strip(),
        status="draft",
        estimated_palettes=estimated_palettes,
        proposed_rate_eur=proposed_rate or None,
        total_eur=total or None,
        valid_until=_date.fromisoformat(valid_until) if valid_until else None,
        notes=notes,
    )
    db.add(offer)
    await db.flush()
    await activity_record(
        db, action="create", user_id=user.id, user_name=user.full_name or user.username,
        user_role=user.role, module="commercial", entity_type="rate_offer",
        entity_id=offer.id, entity_label=ref, ip_address=_client_ip(request),
    )
    return RedirectResponse(url="/commercial/offers", status_code=303)


@router.post("/offers/{offer_id}/send")
async def offer_send(
    offer_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    o = await db.get(RateOffer, offer_id)
    if o is None:
        raise HTTPException(status_code=404)
    o.status = "sent"
    o.sent_at = datetime.now(timezone.utc)
    await db.flush()
    await activity_record(
        db, action="update", user_id=user.id, user_name=user.full_name or user.username,
        user_role=user.role, module="commercial", entity_type="rate_offer",
        entity_id=o.id, entity_label=o.reference, detail="sent",
        ip_address=_client_ip(request),
    )
    return RedirectResponse(url="/commercial/offers", status_code=303)


@router.post("/offers/{offer_id}/convert")
async def offer_convert_to_order(
    offer_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    """Convertir une offre en commande ferme."""
    offer = await db.get(RateOffer, offer_id)
    if offer is None:
        raise HTTPException(status_code=404)
    if offer.status not in ("draft", "sent", "accepted"):
        raise HTTPException(status_code=400, detail="offre non convertible")
    offer.status = "accepted"
    offer.accepted_at = datetime.now(timezone.utc)
    ref = await next_order_reference(db)
    order = Order(
        reference=ref,
        client_id=offer.client_id,
        offer_id=offer.id,
        leg_id=offer.leg_id,
        status="draft",
        booked_palettes=offer.estimated_palettes or 0,
        rate_per_palette_eur=offer.proposed_rate_eur,
        total_eur=offer.total_eur,
        pipedrive_deal_id=offer.pipedrive_deal_id,
    )
    db.add(order)
    await db.flush()
    await activity_record(
        db, action="create", user_id=user.id, user_name=user.full_name or user.username,
        user_role=user.role, module="commercial", entity_type="order",
        entity_id=order.id, entity_label=ref,
        detail=f"from offer {offer.reference}",
        ip_address=_client_ip(request),
    )
    return RedirectResponse(url=f"/commercial/orders/{order.id}", status_code=303)


# ────────────────────────────────────────────── Orders
@router.get("/orders", response_class=HTMLResponse)
async def orders_list(
    request: Request,
    status: str | None = None,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    stmt = select(Order).options(selectinload(Order.client))
    if status:
        stmt = stmt.where(Order.status == status)
    stmt = stmt.order_by(Order.created_at.desc())
    orders = list((await db.execute(stmt)).scalars().all())
    return templates.TemplateResponse(
        "staff/commercial/orders.html",
        {"request": request, "user": user, "orders": orders, "filter_status": status},
    )


@router.get("/orders/new", response_class=HTMLResponse)
async def order_new_form(
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
) -> HTMLResponse:
    clients = list((await db.execute(
        select(Client).where(Client.is_active.is_(True)).order_by(Client.name)
    )).scalars().all())
    legs = list((await db.execute(
        select(Leg).order_by(Leg.etd.desc()).limit(50)
    )).scalars().all())
    return templates.TemplateResponse(
        "staff/commercial/order_form.html",
        {"request": request, "user": user, "clients": clients,
         "legs": legs, "order": None},
    )


@router.post("/orders")
async def order_create(
    request: Request,
    client_id: int = Form(...),
    leg_id: int | None = Form(None),
    booked_palettes: int = Form(0),
    rate_per_palette_eur: float | None = Form(None),
    cargo_description: str | None = Form(None),
    shipper_name: str | None = Form(None),
    consignee_name: str | None = Form(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    if not await db.get(Client, client_id):
        raise HTTPException(status_code=404, detail="client introuvable")
    ref = await next_order_reference(db)
    rate = Decimal(str(rate_per_palette_eur)) if rate_per_palette_eur else None
    total = (rate * Decimal(booked_palettes)).quantize(Decimal("0.01")) if rate else None
    order = Order(
        reference=ref,
        client_id=client_id,
        leg_id=leg_id,
        status="draft",
        booked_palettes=booked_palettes,
        rate_per_palette_eur=rate,
        total_eur=total,
        cargo_description=cargo_description,
        shipper_name=shipper_name,
        consignee_name=consignee_name,
    )
    db.add(order)
    await db.flush()
    await activity_record(
        db, action="create", user_id=user.id, user_name=user.full_name or user.username,
        user_role=user.role, module="commercial", entity_type="order",
        entity_id=order.id, entity_label=ref, ip_address=_client_ip(request),
    )
    return RedirectResponse(url=f"/commercial/orders/{order.id}", status_code=303)


@router.get("/orders/{order_id}", response_class=HTMLResponse)
async def order_detail(
    order_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    order = (await db.execute(
        select(Order).options(selectinload(Order.client), selectinload(Order.assignments))
        .where(Order.id == order_id)
    )).scalar_one_or_none()
    if order is None:
        raise HTTPException(status_code=404)
    leg = await db.get(Leg, order.leg_id) if order.leg_id else None
    return templates.TemplateResponse(
        "staff/commercial/order_detail.html",
        {"request": request, "user": user, "order": order, "leg": leg},
    )


@router.post("/orders/{order_id}/confirm")
async def order_confirm(
    order_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    order = await db.get(Order, order_id)
    if order is None:
        raise HTTPException(status_code=404)
    order.status = "confirmed"
    order.confirmed_at = datetime.now(timezone.utc)
    await db.flush()
    await activity_record(
        db, action="update", user_id=user.id, user_name=user.full_name or user.username,
        user_role=user.role, module="commercial", entity_type="order",
        entity_id=order.id, entity_label=order.reference, detail="confirmed",
        ip_address=_client_ip(request),
    )
    return RedirectResponse(url=f"/commercial/orders/{order_id}", status_code=303)


@router.post("/orders/{order_id}/cancel")
async def order_cancel(
    order_id: int,
    request: Request,
    reason: str = Form(""),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "S")),
):
    order = await db.get(Order, order_id)
    if order is None:
        raise HTTPException(status_code=404)
    order.status = "cancelled"
    order.cancelled_at = datetime.now(timezone.utc)
    order.cancelled_reason = reason or None
    await db.flush()
    await activity_record(
        db, action="delete", user_id=user.id, user_name=user.full_name or user.username,
        user_role=user.role, module="commercial", entity_type="order",
        entity_id=order.id, entity_label=order.reference,
        detail=f"cancelled: {reason}", ip_address=_client_ip(request),
    )
    return RedirectResponse(url="/commercial/orders", status_code=303)


# ────────────────────────────────────────────── Offer DOCX export
@router.get("/offers/{offer_id}/export.docx")
async def offer_export_docx(
    offer_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> Response:
    """Exporte une offre commerciale en fichier Word (.docx)."""
    offer = await db.get(RateOffer, offer_id)
    if offer is None:
        raise HTTPException(status_code=404, detail="offre introuvable")

    client = await db.get(Client, offer.client_id)
    leg = await db.get(Leg, offer.leg_id) if offer.leg_id else None

    # ── Build document ──────────────────────────────────────────────────────
    doc = Document()

    # ── Header: OFFRE COMMERCIALE + reference
    heading = doc.add_heading("", level=0)
    run = heading.add_run("OFFRE COMMERCIALE NEWTOWT")
    run.font.color.rgb = RGBColor(0x0D, 0x59, 0x66)
    run.font.size = Pt(20)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref_run = ref_para.add_run(f"Référence : {offer.reference}")
    ref_run.bold = True
    ref_run.font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # ── Section Client ──────────────────────────────────────────────────────
    client_heading = doc.add_heading("Client", level=2)
    client_heading.runs[0].font.color.rgb = RGBColor(0x0D, 0x59, 0x66)

    client_table = doc.add_table(rows=0, cols=2)
    client_table.style = "Table Grid"

    def _add_kv_row(table, label: str, value: str) -> None:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value

    _add_kv_row(client_table, "Nom", client.name if client else "—")
    if client and client.company_name:
        _add_kv_row(client_table, "Société", client.company_name)
    _add_kv_row(client_table, "E-mail", client.email if client else "—")
    _add_kv_row(client_table, "Téléphone", client.phone if client else "—")

    doc.add_paragraph()  # spacer

    # ── Section Objet ───────────────────────────────────────────────────────
    objet_heading = doc.add_heading("Objet", level=2)
    objet_heading.runs[0].font.color.rgb = RGBColor(0x0D, 0x59, 0x66)
    doc.add_paragraph(offer.title or "—")

    doc.add_paragraph()  # spacer

    # ── Section Itinéraire ──────────────────────────────────────────────────
    itin_heading = doc.add_heading("Itinéraire", level=2)
    itin_heading.runs[0].font.color.rgb = RGBColor(0x0D, 0x59, 0x66)

    if leg:
        etd_str = leg.etd.strftime("%d/%m/%Y") if leg.etd else "—"
        eta_str = leg.eta.strftime("%d/%m/%Y") if leg.eta else "—"
        doc.add_paragraph(
            f"Leg : {leg.leg_code}\n"
            f"ETD : {etd_str}     ETA : {eta_str}"
        )
    else:
        doc.add_paragraph("À confirmer")

    doc.add_paragraph()  # spacer

    # ── Table Tarification ──────────────────────────────────────────────────
    tarif_heading = doc.add_heading("Tarification", level=2)
    tarif_heading.runs[0].font.color.rgb = RGBColor(0x0D, 0x59, 0x66)

    tarif_table = doc.add_table(rows=1, cols=4)
    tarif_table.style = "Table Grid"
    hdr_cells = tarif_table.rows[0].cells
    for idx, label in enumerate(["Description", "Quantité", "Tarif unitaire", "Total"]):
        hdr_cells[idx].text = label
        hdr_cells[idx].paragraphs[0].runs[0].bold = True

    row_cells = tarif_table.add_row().cells
    qty = offer.estimated_palettes or 0
    rate = offer.proposed_rate_eur
    total = offer.total_eur
    rate_str = f"{rate:,.2f} EUR/palette".replace(",", " ") if rate is not None else "—"
    total_str = f"{total:,.2f} EUR".replace(",", " ") if total is not None else "—"
    row_cells[0].text = "Fret aérien palettes"
    row_cells[1].text = str(qty)
    row_cells[2].text = rate_str
    row_cells[3].text = total_str

    doc.add_paragraph()  # spacer

    # ── Section Conditions ──────────────────────────────────────────────────
    cond_heading = doc.add_heading("Conditions", level=2)
    cond_heading.runs[0].font.color.rgb = RGBColor(0x0D, 0x59, 0x66)

    validity_str = (
        offer.valid_until.strftime("%d/%m/%Y") if offer.valid_until else "—"
    )
    cond_para = doc.add_paragraph()
    cond_para.add_run(f"Validité : {validity_str}\n").bold = False
    cond_para.add_run(
        "Ce prix inclut le transport par voilier cargo à propulsion vélique "
        "(zéro émission directe)."
    )

    # ── Section Notes (optionnelle) ─────────────────────────────────────────
    if offer.notes:
        doc.add_paragraph()  # spacer
        notes_heading = doc.add_heading("Notes", level=2)
        notes_heading.runs[0].font.color.rgb = RGBColor(0x0D, 0x59, 0x66)
        doc.add_paragraph(offer.notes)

    doc.add_paragraph()  # spacer

    # ── Footer ──────────────────────────────────────────────────────────────
    footer_para = doc.add_paragraph(
        "NEWTOWT — Pioneer of wind-powered cargo since 2011 — www.newtowt.eu"
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x59, 0x66)
        run.font.size = Pt(9)
        run.italic = True

    # ── Serialize ───────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    await activity_record(
        db, action="offer_export_docx",
        user_id=user.id, user_name=user.full_name or user.username,
        user_role=user.role, module="commercial", entity_type="rate_offer",
        entity_id=offer.id, entity_label=offer.reference,
        detail=offer.reference,
        ip_address=_client_ip(request),
    )

    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="Offre_{offer.reference}.docx"'
        },
    )


def _client_ip(request: Request) -> str | None:
    return request.headers.get("x-forwarded-for") or (request.client.host if request.client else None)
